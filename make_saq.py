from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# ── Styles ───────────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def numbered(items, start=1):
    for i, item in enumerate(items, start):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(item)
        run.font.size = Pt(11)

def lettered(items):
    for letter, item in zip('abcdefgh', items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"({letter})  {item}")
        run.font.size = Pt(11)

# ── Title ─────────────────────────────────────────────────────────────────────
t = doc.add_heading("STANDARD ASSESSMENT QUESTIONS", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Unit: BS/C/7119 — Apply Fundamentals of Accounting\nWeek 5: Account for Fixed Assets (Non-Current Assets)")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.size = Pt(11)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
heading("SECTION A: SHORT-ANSWER AND COMPUTATION QUESTIONS")
doc.add_paragraph()

# ── Q1 ───────────────────────────────────────────────────────────────────────
heading("Question 1 — Cost of a Fixed Asset  (LO2)", level=2)
body("Mwangi is setting up a factory in Athi River. He incurs the following costs in acquiring a new machine:")
doc.add_paragraph()

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
headers = ["Item", "Amount (KES)"]
rows_data = [
    ("Purchase price", "800,000"),
    ("Delivery charges", "25,000"),
    ("Installation costs", "40,000"),
    ("Testing and trial runs", "15,000"),
]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
for r, (item, amt) in enumerate(rows_data, 1):
    tbl.rows[r].cells[0].text = item
    tbl.rows[r].cells[1].text = amt

doc.add_paragraph()
body("Required:")
lettered(["Calculate the TOTAL COST to be capitalised for this machine."])
doc.add_paragraph()

# ── Q2 ───────────────────────────────────────────────────────────────────────
heading("Question 2 — Straight-Line Depreciation  (LO3)", level=2)
body("Wanjiku Transport, Nairobi, purchases a delivery van with the following details:")
lettered([
    "Cost: KES 2,000,000",
    "Estimated residual value: KES 200,000",
    "Estimated useful life: 5 years",
    "Depreciation method: Straight-Line",
])
doc.add_paragraph()
body("Required:")
lettered([
    "Calculate the annual depreciation charge.",
    "Prepare a depreciation schedule for all 5 years.",
    "State the journal entry to record depreciation each year.",
])
doc.add_paragraph()

# ── Q3 ───────────────────────────────────────────────────────────────────────
heading("Question 3 — Reducing Balance Depreciation  (LO3)", level=2)
body("Ochieng Engineering, Kisumu, purchases machinery with the following details:")
lettered([
    "Cost: KES 1,000,000",
    "Depreciation rate: 25% per annum (reducing balance)",
])
doc.add_paragraph()
body("Required:")
lettered([
    "Calculate the depreciation charge for each of the first 5 years.",
    "Prepare a depreciation schedule showing the Net Book Value at the end of each year.",
    "State the journal entry to record depreciation in Year 1.",
])
doc.add_paragraph()

# ── Q4 ───────────────────────────────────────────────────────────────────────
heading("Question 4 — Provision for Depreciation Ledger  (LO4 & LO5)", level=2)
body("Using Wanjiku Transport's delivery van (Cost: KES 2,000,000; Annual depreciation: KES 360,000 straight-line):")
doc.add_paragraph()
body("Required — prepare the following ledger accounts for Years 1, 2, and 3:")
lettered([
    "Motor Vehicle Account (at cost)",
    "Provision for Depreciation — Motor Vehicle Account",
    "Depreciation Expense Account",
])
doc.add_paragraph()

# ── Q5 ───────────────────────────────────────────────────────────────────────
heading("Question 5 — Disposal at a Profit  (LO6 & LO7)", level=2)
body("Kamau Traders, Nakuru, purchased office furniture on 1 January 2023 for KES 200,000. "
     "Depreciation policy: 10% per annum, straight-line method. "
     "On 1 January 2026, the furniture was sold for KES 160,000 cash.")
doc.add_paragraph()
body("Required:")
lettered([
    "Calculate the accumulated depreciation at the date of disposal.",
    "Calculate the Net Book Value at the date of disposal.",
    "Determine whether there is a profit or loss on disposal, and state the amount.",
    "Prepare the Disposal Account.",
    "Record all necessary journal entries.",
])
doc.add_paragraph()

# ── Q6 ───────────────────────────────────────────────────────────────────────
heading("Question 6 — Disposal at a Loss  (LO6 & LO7)", level=2)
body("Atieno Supplies, Kisumu, purchased a computer on 1 January 2024 for KES 120,000. "
     "Depreciation policy: 25% per annum, reducing balance method. "
     "On 31 December 2025, the computer was sold for KES 50,000.")
doc.add_paragraph()
body("Required:")
lettered([
    "Calculate the depreciation charge for Year 1 and Year 2.",
    "Calculate the accumulated depreciation and Net Book Value at the date of disposal.",
    "Determine whether there is a profit or loss on disposal, and state the amount.",
    "Prepare the Disposal Account.",
    "Record all necessary journal entries.",
])
doc.add_paragraph()

# ── Q7 ───────────────────────────────────────────────────────────────────────
heading("Question 7 — Depreciation Schedules  (LO3)", level=2)
body("Nyambura Agribusiness, Nanyuki, purchased a tractor on 1 July 2024 for KES 3,600,000. "
     "The estimated residual value is KES 600,000 and the useful life is 5 years. "
     "The company's financial year ends on 30 June.")
doc.add_paragraph()
body("Required:")
lettered([
    "Prepare a depreciation schedule for the full 5 years using the Straight-Line method.",
    "Prepare a depreciation schedule for the full 5 years using the Reducing Balance method at 30% per annum.",
])
doc.add_paragraph()

# ── Q8 ───────────────────────────────────────────────────────────────────────
heading("Question 8 — Disposal with Time Apportionment  (LO6 & LO7)", level=2)
body("Mutua Logistics, Machakos, purchased a lorry on 1 January 2023 for KES 4,000,000. "
     "Depreciation is charged at 20% per annum using the straight-line method with no residual value. "
     "On 30 June 2025, the lorry was sold for KES 2,200,000.")
doc.add_paragraph()
body("Required:")
lettered([
    "Calculate the accumulated depreciation at the date of disposal.",
    "Prepare the Disposal Account.",
    "Record all necessary journal entries.",
    "State whether there was a profit or loss on disposal and show its treatment in the financial statements.",
])
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
heading("SECTION B: DISCUSSION QUESTIONS")
doc.add_paragraph()

heading("Question 9  (LO1)", level=2)
body("(a)  Define the term 'fixed asset' and state TWO characteristics that distinguish a fixed asset from a current asset.")
body("(b)  State the TWO conditions that must be met before a fixed asset can be recognised in the books of account under IAS 16.")
doc.add_paragraph()

heading("Question 10  (LO3)", level=2)
body("(a)  Explain why land is generally not depreciated but buildings are.")
body("(b)  A business in Nairobi purchased computers in 2020. Explain why the reducing balance method may be more appropriate than the straight-line method for depreciating these computers.")
doc.add_paragraph()

heading("Question 11  (LO8)", level=2)
body("(a)  State TWO accounting principles that support the practice of depreciating fixed assets. For each principle, explain how it applies to depreciation.")
body("(b)  Explain the Historical Cost principle and state how it applies to the initial recording of fixed assets.")
doc.add_paragraph()

heading("Question 12  (LO2)", level=2)
body("A business purchases a machine for KES 500,000. In addition, it pays KES 20,000 for delivery, KES 15,000 for installation, and KES 5,000 for a maintenance contract for the first year.")
body("(a)  State the total cost to be capitalised for this machine.")
body("(b)  Explain why the maintenance contract cost should NOT be capitalised.")
doc.add_paragraph()

# ── Save ─────────────────────────────────────────────────────────────────────
out = "/home/stdk/NOTES_KSTVET/Fundamentals of Accountig _BS_C_7119/Standard Assessment Questions.docx"
doc.save(out)
print(f"Saved: {out}")
